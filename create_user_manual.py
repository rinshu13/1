from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "Универсальный_Руководство_пользователя.docx"


def set_cell_text(cell, text, bold=False, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Оглавление обновляется в Word: выделите поле и нажмите F9."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def set_paragraph_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_para(doc, text="", style=None, align=None):
    p = doc.add_paragraph(text, style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_font(p)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    set_paragraph_font(p)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    set_paragraph_font(p)
    return p


def add_placeholder(doc, caption, height_rows=4):
    table = doc.add_table(rows=height_rows, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        row.height = Cm(1.2)
        cell = row.cells[0]
        set_cell_width(cell, 16)
        shade_cell(cell, "F2F2F2")
    merged = table.cell(0, 0)
    for i in range(1, height_rows):
        merged = merged.merge(table.cell(i, 0))
    set_cell_text(merged, caption, bold=True, size=14)
    p = doc.add_paragraph(caption.replace("МЕСТО ДЛЯ ", "").capitalize())
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(p, 12)
    return table


doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(1.5)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
styles["Normal"].font.size = Pt(14)
styles["Normal"].paragraph_format.line_spacing = 1.5
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size in [("Heading 1", 16), ("Heading 2", 15), ("Heading 3", 14)]:
    style = styles[name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

# Title page
for _ in range(2):
    doc.add_paragraph()

p = add_para(doc, "[ПОЛНОЕ НАИМЕНОВАНИЕ ОБРАЗОВАТЕЛЬНОЙ ОРГАНИЗАЦИИ]", align=WD_ALIGN_PARAGRAPH.CENTER)
p.runs[0].bold = True

for _ in range(5):
    doc.add_paragraph()

p = add_para(doc, "РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ", align=WD_ALIGN_PARAGRAPH.CENTER)
p.runs[0].bold = True
p.runs[0].font.size = Pt(18)
p = add_para(doc, "к программному продукту", align=WD_ALIGN_PARAGRAPH.CENTER)
p.runs[0].bold = True
p = add_para(doc, "[НАЗВАНИЕ ПРОГРАММНОГО ПРОДУКТА]", align=WD_ALIGN_PARAGRAPH.CENTER)
p.runs[0].bold = True

for _ in range(6):
    doc.add_paragraph()

p = add_para(doc, "Разработчик: [ФИО обучающегося]", align=WD_ALIGN_PARAGRAPH.RIGHT)
p = add_para(doc, "Группа: [номер группы]", align=WD_ALIGN_PARAGRAPH.RIGHT)
p = add_para(doc, "Специальность: 09.02.07 Информационные системы и программирование", align=WD_ALIGN_PARAGRAPH.RIGHT)

for _ in range(5):
    doc.add_paragraph()

add_para(doc, "[ГОРОД] - [ГОД]", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

doc.add_heading("Содержание", level=1)
add_toc(doc.add_paragraph())
doc.add_page_break()

doc.add_heading("1 Основные сведения о программном продукте", level=1)
add_para(doc, "Программный продукт [НАЗВАНИЕ ПРОГРАММНОГО ПРОДУКТА] предназначен для автоматизации основных операций организации в выбранной предметной области. Система обеспечивает хранение, просмотр, поиск, добавление, изменение и удаление записей, а также разграничение доступных действий в зависимости от роли пользователя.")
add_para(doc, "Предметная область системы: [УКАЖИТЕ ПРЕДМЕТНУЮ ОБЛАСТЬ, НАПРИМЕР МАГАЗИН, САЛОН, СКЛАД, СЕРВИСНЫЙ ЦЕНТР].")
add_para(doc, "Основными пользователями системы являются сотрудники организации, выполняющие учет, обработку и контроль данных.")

doc.add_heading("1.1 Назначение системы", level=2)
add_para(doc, "Система позволяет выполнять следующие действия:")
for item in [
    "вход в систему под учетной записью пользователя;",
    "просмотр списка основных объектов предметной области;",
    "поиск записей по ключевым значениям;",
    "фильтрация и сортировка отображаемых данных;",
    "добавление новых записей;",
    "изменение существующих записей;",
    "удаление или деактивация записей;",
    "просмотр информации о скидках, остатках и других важных показателях."
]:
    add_bullet(doc, item)

doc.add_heading("1.2 Роли пользователей", level=2)
add_para(doc, "Перечень ролей может быть изменен в соответствии с заданием. В универсальном варианте предусмотрены следующие роли.")
roles = doc.add_table(rows=1, cols=3)
roles.style = "Table Grid"
roles.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Роль", "Назначение", "Основные возможности"]
for i, h in enumerate(headers):
    set_cell_text(roles.rows[0].cells[i], h, bold=True)
    shade_cell(roles.rows[0].cells[i], "E8EEF5")
for role, purpose, actions in [
    ("Администратор", "Полное сопровождение системы", "Работа с пользователями, справочниками и всеми данными системы."),
    ("Сотрудник", "Ежедневная работа с данными", "Просмотр, поиск, добавление и изменение записей в пределах полномочий."),
    ("Пользователь", "Просмотр информации", "Просмотр списка записей и получение справочной информации.")
]:
    row = roles.add_row().cells
    set_cell_text(row[0], role)
    set_cell_text(row[1], purpose)
    set_cell_text(row[2], actions)
add_para(doc, "Таблица 1 - Роли пользователей системы", align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("1.3 Требования к запуску", level=2)
add_para(doc, "Для работы с программным продуктом требуется компьютер с операционной системой Windows, установленной платформой .NET Framework и доступом к базе данных. Перед началом работы база данных должна быть создана и подключена к приложению.")

doc.add_heading("2 Начало работы с системой", level=1)
doc.add_heading("2.1 Запуск приложения", level=2)
add_para(doc, "Для запуска программы пользователь открывает исполняемый файл приложения или запускает проект из среды разработки. После запуска отображается форма авторизации.")
add_placeholder(doc, "МЕСТО ДЛЯ СКРИНШОТА: форма авторизации пользователя")

doc.add_heading("2.2 Авторизация пользователя", level=2)
add_para(doc, "Для входа в систему необходимо выполнить следующие действия:")
for item in [
    "ввести логин в поле «Логин»;",
    "ввести пароль в поле «Пароль»;",
    "нажать кнопку «Войти»;",
    "при успешной проверке учетных данных перейти к главной форме приложения."
]:
    add_number(doc, item)
add_para(doc, "Если логин или пароль указаны неверно, система выводит сообщение об ошибке. В этом случае необходимо проверить правильность введенных данных и повторить попытку.")

doc.add_heading("3 Главное окно приложения", level=1)
add_para(doc, "Главное окно содержит заголовок программы, логотип, панель управляющих кнопок, область поиска и фильтрации, а также список записей. Внешний вид главного окна должен соответствовать руководству по стилю, указанному в задании.")
add_placeholder(doc, "МЕСТО ДЛЯ СКРИНШОТА: главное окно приложения")

doc.add_heading("3.1 Список записей", level=2)
add_para(doc, "Список записей отображает основные данные из базы данных. В зависимости от предметной области это могут быть товары, услуги, клиенты, заказы, заявки или другие объекты. Каждая запись содержит ключевую информацию, необходимую пользователю для принятия решения.")
add_para(doc, "Если в задании предусмотрено карточное отображение, каждая запись может включать изображение, наименование, категорию, описание, цену, скидку и количество на складе. Если используется табличное отображение, данные выводятся в виде строк и столбцов.")
add_placeholder(doc, "МЕСТО ДЛЯ СКРИНШОТА: список записей или карточки объектов")

doc.add_heading("3.2 Поиск, фильтрация и сортировка", level=2)
add_para(doc, "Для быстрого нахождения нужной записи пользователь может использовать поиск, фильтр и сортировку.")
for item in [
    "поиск выполняется по названию, артикулу, описанию или другому ключевому полю;",
    "фильтр ограничивает список выбранной категорией, типом или другим справочным значением;",
    "сортировка изменяет порядок записей по названию, цене, дате или иному показателю."
]:
    add_bullet(doc, item)
add_placeholder(doc, "МЕСТО ДЛЯ СКРИНШОТА: поиск и фильтрация данных")

doc.add_heading("4 Работа с записями", level=1)
doc.add_heading("4.1 Добавление записи", level=2)
add_para(doc, "Для добавления новой записи пользователь нажимает кнопку «Добавить». После этого открывается форма ввода данных.")
add_para(doc, "Пользователь заполняет обязательные поля, проверяет корректность введенной информации и нажимает кнопку «Сохранить». После сохранения новая запись появляется в общем списке.")
add_placeholder(doc, "МЕСТО ДЛЯ СКРИНШОТА: форма добавления записи")

doc.add_heading("4.2 Изменение записи", level=2)
add_para(doc, "Для изменения существующей записи необходимо выбрать ее в списке и нажать кнопку «Изменить» либо дважды щелкнуть по записи. Система открывает форму редактирования с уже заполненными полями.")
add_para(doc, "После внесения изменений пользователь нажимает кнопку «Сохранить». Если изменения выполнены корректно, список обновляется.")
add_placeholder(doc, "МЕСТО ДЛЯ СКРИНШОТА: форма изменения записи")

doc.add_heading("4.3 Удаление записи", level=2)
add_para(doc, "Для удаления записи пользователь выбирает ее в списке и нажимает кнопку «Удалить». Перед выполнением действия система запрашивает подтверждение. После подтверждения запись удаляется из отображаемого списка или помечается как неактивная.")
add_para(doc, "Если запись связана с другими данными, рекомендуется использовать деактивацию вместо полного удаления, чтобы не нарушить целостность базы данных.")
add_placeholder(doc, "МЕСТО ДЛЯ СКРИНШОТА: подтверждение удаления записи")

doc.add_heading("5 Описание функционала по ролям", level=1)
doc.add_heading("5.1 Администратор", level=2)
add_para(doc, "Администратор имеет доступ ко всем основным разделам системы. Он может управлять пользователями, справочниками, основными записями и настройками приложения.")
doc.add_heading("5.2 Сотрудник", level=2)
add_para(doc, "Сотрудник выполняет повседневные операции: просматривает список, ищет нужные записи, добавляет новые данные и редактирует информацию в рамках выданных прав.")
doc.add_heading("5.3 Пользователь с ограниченными правами", level=2)
add_para(doc, "Пользователь с ограниченными правами может просматривать информацию и использовать поиск, фильтрацию и сортировку. Изменение данных для этой роли может быть недоступно.")

doc.add_heading("6 Особенности отображения данных", level=1)
add_para(doc, "При отображении данных система может выделять записи цветом, если требуется обратить внимание пользователя на важное состояние объекта.")
for item in [
    "если размер скидки превышает установленное значение, запись выделяется специальным цветом;",
    "если цена снижена, исходная цена может отображаться перечеркнутой, а итоговая цена указывается рядом;",
    "если объект отсутствует на складе или недоступен, запись выделяется отдельным цветом;",
    "если данные заполнены некорректно, система должна вывести понятное сообщение пользователю."
]:
    add_bullet(doc, item)
add_placeholder(doc, "МЕСТО ДЛЯ СКРИНШОТА: выделение записей по условию")

doc.add_heading("7 Схема базы данных", level=1)
add_para(doc, "В данном разделе размещается схема базы данных. На схеме должны быть показаны основные таблицы, ключевые поля и связи между таблицами.")
add_placeholder(doc, "МЕСТО ДЛЯ СХЕМЫ: схема базы данных", height_rows=6)
add_para(doc, "Рисунок 8 - Схема базы данных", align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("8 Возможные ошибки и способы их устранения", level=1)
errors = doc.add_table(rows=1, cols=3)
errors.style = "Table Grid"
errors.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Ситуация", "Причина", "Действия пользователя"]):
    set_cell_text(errors.rows[0].cells[i], h, bold=True)
    shade_cell(errors.rows[0].cells[i], "E8EEF5")
for situation, reason, action in [
    ("Не выполняется вход", "Неверный логин или пароль", "Проверить введенные данные и повторить вход."),
    ("Данные не отображаются", "Нет подключения к базе данных", "Проверить запуск сервера базы данных и настройки подключения."),
    ("Запись не сохраняется", "Не заполнены обязательные поля", "Заполнить поля, отмеченные как обязательные, и повторить сохранение."),
    ("Запись не удаляется", "Запись связана с другими данными", "Использовать деактивацию записи или обратиться к администратору.")
]:
    row = errors.add_row().cells
    set_cell_text(row[0], situation)
    set_cell_text(row[1], reason)
    set_cell_text(row[2], action)
add_para(doc, "Таблица 2 - Возможные ошибки и способы их устранения", align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("9 Завершение работы", level=1)
add_para(doc, "Для завершения работы пользователь закрывает главное окно приложения стандартной кнопкой закрытия окна. Если в системе предусмотрена кнопка выхода, пользователь нажимает ее и подтверждает завершение работы.")

doc.add_heading("Приложение А. Перечень иллюстраций", level=1)
for item in [
    "Рисунок 1 - Форма авторизации пользователя",
    "Рисунок 2 - Главное окно приложения",
    "Рисунок 3 - Список записей или карточки объектов",
    "Рисунок 4 - Поиск и фильтрация данных",
    "Рисунок 5 - Форма добавления записи",
    "Рисунок 6 - Форма изменения записи",
    "Рисунок 7 - Подтверждение удаления записи",
    "Рисунок 8 - Схема базы данных"
]:
    add_bullet(doc, item)

doc.add_heading("Приложение Б. Памятка по заполнению шаблона", level=1)
for item in [
    "Замените все фрагменты в квадратных скобках на данные своей предметной области.",
    "Вставьте реальные скриншоты вместо серых блоков-заполнителей.",
    "После вставки скриншотов обновите номера страниц и оглавление.",
    "Проверьте, что описание ролей совпадает с реализованными возможностями приложения.",
    "Удалите этот раздел перед сдачей, если он не требуется по заданию."
]:
    add_bullet(doc, item)

doc.save(OUT)
print(OUT)
